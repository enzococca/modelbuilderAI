"""Workflow CRUD and execution API endpoints."""

from __future__ import annotations

import asyncio
import csv as csv_mod
import io
import json
import os
import re
import uuid
import zipfile

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from models.workflow_models import WorkflowCreate, WorkflowOut
from orchestrator.pipeline import PipelineExecutor
from storage.database import get_db, WorkflowRow, WorkflowExecutionRow

router = APIRouter(tags=["workflows"])


@router.get("/workflows")
async def list_workflows(db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(WorkflowRow).order_by(WorkflowRow.created_at.desc()))
    rows = result.scalars().all()
    return [
        WorkflowOut(
            id=r.id, name=r.name, description=r.description,
            definition=r.definition or {"nodes": [], "edges": []},
            created_at=r.created_at.isoformat(),
            updated_at=r.updated_at.isoformat() if r.updated_at else r.created_at.isoformat(),
        )
        for r in rows
    ]


@router.post("/workflows", status_code=201)
async def create_workflow(body: WorkflowCreate, db: AsyncSession = Depends(get_db)):
    row = WorkflowRow(
        id=str(uuid.uuid4()),
        name=body.name,
        description=body.description,
        definition=body.definition.model_dump(),
    )
    db.add(row)
    await db.commit()
    return WorkflowOut(
        id=row.id, name=row.name, description=row.description,
        definition=body.definition,
        created_at=row.created_at.isoformat(),
        updated_at=row.created_at.isoformat(),
    )


@router.get("/workflows/{workflow_id}")
async def get_workflow(workflow_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Workflow not found")
    return WorkflowOut(
        id=row.id, name=row.name, description=row.description,
        definition=row.definition or {"nodes": [], "edges": []},
        created_at=row.created_at.isoformat(),
        updated_at=row.updated_at.isoformat() if row.updated_at else row.created_at.isoformat(),
    )


@router.put("/workflows/{workflow_id}")
async def update_workflow(workflow_id: str, body: WorkflowCreate, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Workflow not found")
    row.name = body.name
    row.description = body.description
    row.definition = body.definition.model_dump()
    await db.commit()
    return WorkflowOut(
        id=row.id, name=row.name, description=row.description,
        definition=body.definition,
        created_at=row.created_at.isoformat(),
        updated_at=row.updated_at.isoformat() if row.updated_at else row.created_at.isoformat(),
    )


@router.delete("/workflows/{workflow_id}", status_code=204)
async def delete_workflow(workflow_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Workflow not found")
    await db.delete(row)
    await db.commit()


@router.post("/workflows/{workflow_id}/validate")
async def validate_workflow(workflow_id: str, db: AsyncSession = Depends(get_db)):
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Workflow not found")
    from models.workflow_models import WorkflowDefinition
    defn = WorkflowDefinition(**(row.definition or {"nodes": [], "edges": []}))
    executor = PipelineExecutor(defn)
    errors = executor.validate()
    return {"valid": len(errors) == 0, "errors": errors}


@router.post("/workflows/{workflow_id}/run")
async def run_workflow(workflow_id: str, db: AsyncSession = Depends(get_db)):
    """Launch workflow execution as a background task with WebSocket updates."""
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    row = result.scalar_one_or_none()
    if not row:
        raise HTTPException(404, "Workflow not found")
    from models.workflow_models import WorkflowDefinition
    from builder.workflow_engine import WorkflowEngine
    from websocket.handlers import manager

    defn = WorkflowDefinition(**(row.definition or {"nodes": [], "edges": []}))
    engine = WorkflowEngine(defn, workflow_id=workflow_id, broadcast=manager.broadcast)

    async def _run_in_background() -> None:
        from storage.database import async_session as get_session
        try:
            results = await engine.run()
            # Persist execution results
            async with get_session() as session:
                session.add(WorkflowExecutionRow(
                    workflow_id=workflow_id,
                    status="completed",
                    results={k: str(v) for k, v in results.items()},
                ))
                await session.commit()
        except Exception as e:
            async with get_session() as session:
                session.add(WorkflowExecutionRow(
                    workflow_id=workflow_id,
                    status="error",
                    error=str(e),
                ))
                await session.commit()
            await manager.broadcast({
                "type": "workflow_status",
                "workflow_id": workflow_id,
                "status": "error",
                "node_statuses": {},
                "results": {},
                "error": str(e),
            })

    asyncio.create_task(_run_in_background())
    return {"status": "started", "workflow_id": workflow_id}


def _parse_markdown_tables(text: str) -> list[list[str]]:
    """Extract markdown tables from text as list of rows (each row = list of cells)."""
    rows: list[list[str]] = []
    for line in text.splitlines():
        line = line.strip()
        if not line.startswith("|") or not line.endswith("|"):
            continue
        # Skip separator lines (e.g. |---|---|)
        cells = [c.strip() for c in line.strip("|").split("|")]
        if all(re.match(r"^[-:]+$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def _md_to_plain(text: str) -> str:
    """Strip common markdown formatting for plain-text export."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
    text = re.sub(r"\*(.+?)\*", r"\1", text)       # italic
    text = re.sub(r"`(.+?)`", r"\1", text)         # inline code
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)  # headings
    return text


async def _smart_format(content: str, target_format: str) -> str:
    """Use a lightweight AI agent to reformat content for the target format."""
    try:
        from orchestrator.router import create_agent

        prompts = {
            "pdf": (
                "Reformat the following text for a professional PDF document. "
                "Remove all markdown syntax (**, *, #, ```, |---|, etc.). "
                "Convert markdown headings to UPPERCASE titles. "
                "Convert markdown lists to plain numbered/bulleted text. "
                "Remove markdown table pipes and format as aligned text. "
                "Keep the content intact, just clean the formatting. "
                "Return ONLY the reformatted text, no explanations."
            ),
            "docx": (
                "Reformat the following text for a Microsoft Word document. "
                "Keep markdown headings (# lines) as they'll become Word headings. "
                "Keep **bold** markers. Convert tables to clean aligned format. "
                "Remove code fences (```). Clean up any raw formatting artifacts. "
                "Return ONLY the reformatted text, no explanations."
            ),
        }
        prompt = prompts.get(target_format)
        if not prompt:
            return content

        agent = create_agent(
            "claude-haiku-4-5-20251001",
            system_prompt=prompt,
            temperature=0.1,
            max_tokens=8192,
        )
        result = await agent.chat([{"role": "user", "content": content}])
        return result if isinstance(result, str) else content
    except Exception:
        return content


@router.post("/workflows/{workflow_id}/export")
async def export_workflow(
    workflow_id: str,
    format: str = Query("zip", pattern="^(zip|zip_all|markdown|pdf|docx|csv|xlsx|png|geojson|shapefile)$"),
    smart_format: bool = Query(False),
    db: AsyncSession = Depends(get_db),
):
    """Export the latest workflow execution results in various formats."""
    # Get workflow info
    result = await db.execute(select(WorkflowRow).where(WorkflowRow.id == workflow_id))
    wf = result.scalar_one_or_none()
    if not wf:
        raise HTTPException(404, "Workflow not found")

    # Get latest execution
    result = await db.execute(
        select(WorkflowExecutionRow)
        .where(WorkflowExecutionRow.workflow_id == workflow_id)
        .order_by(WorkflowExecutionRow.created_at.desc())
        .limit(1)
    )
    execution = result.scalar_one_or_none()
    if not execution or not execution.results:
        raise HTTPException(404, "No execution results found")

    results = execution.results
    defn = wf.definition or {"nodes": [], "edges": []}
    node_labels = {n["id"]: n.get("data", {}).get("label", n["id"]) for n in defn.get("nodes", [])}

    # ── Markdown ──────────────────────────────────────────
    if format == "markdown":
        lines = [f"# {wf.name} — Risultati\n"]
        for nid, content in results.items():
            label = node_labels.get(nid, nid)
            lines.append(f"## {label}\n\n{content}\n")
        md = "\n".join(lines)
        return StreamingResponse(
            io.BytesIO(md.encode()),
            media_type="text/markdown",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.md"'},
        )

    # ── PDF ────────────────────────────────────────────────
    if format == "pdf":
        from fpdf import FPDF

        # Optionally smart-format content
        if smart_format:
            formatted: dict[str, str] = {}
            for nid, content in results.items():
                formatted[nid] = await _smart_format(str(content), "pdf")
        else:
            formatted = {k: str(v) for k, v in results.items()}

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin

        def _safe(text: str) -> str:
            return text.encode("latin-1", "replace").decode("latin-1")

        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(usable_w, 10, _safe(wf.name), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        for nid, content in formatted.items():
            label = node_labels.get(nid, nid)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_x(pdf.l_margin)
            pdf.cell(usable_w, 8, _safe(label), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 10)
            plain = _md_to_plain(content) if not smart_format else content
            for line in plain.splitlines():
                if not line.strip():
                    pdf.ln(3)
                else:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(usable_w, 5, _safe(line))
            pdf.ln(6)

        buf = io.BytesIO(pdf.output())
        return StreamingResponse(
            buf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.pdf"'},
        )

    # ── DOCX ───────────────────────────────────────────────
    if format == "docx":
        from docx import Document
        from docx.shared import Pt

        # Optionally smart-format content
        if smart_format:
            docx_formatted: dict[str, str] = {}
            for nid, content in results.items():
                docx_formatted[nid] = await _smart_format(str(content), "docx")
        else:
            docx_formatted = {k: str(v) for k, v in results.items()}

        doc = Document()
        doc.add_heading(wf.name, level=0)

        for nid, content in docx_formatted.items():
            label = node_labels.get(nid, nid)
            doc.add_heading(label, level=1)
            text = str(content)

            # Check for markdown tables
            tables = _parse_markdown_tables(text)
            if tables:
                # Add table
                tbl = doc.add_table(rows=len(tables), cols=len(tables[0]) if tables else 1)
                tbl.style = "Table Grid"
                for i, row in enumerate(tables):
                    for j, cell in enumerate(row):
                        if j < len(tbl.columns):
                            tbl.rows[i].cells[j].text = cell
                # Add remaining non-table text
                non_table = re.sub(r"\|[^\n]+\|", "", text).strip()
                if non_table:
                    for para_text in non_table.split("\n\n"):
                        if para_text.strip():
                            p = doc.add_paragraph()
                            # Handle bold markdown
                            parts = re.split(r"(\*\*.*?\*\*)", para_text)
                            for part in parts:
                                if part.startswith("**") and part.endswith("**"):
                                    run = p.add_run(part[2:-2])
                                    run.bold = True
                                else:
                                    p.add_run(part)
            else:
                # Plain paragraphs with bold support
                for para_text in text.split("\n\n"):
                    if not para_text.strip():
                        continue
                    p = doc.add_paragraph()
                    parts = re.split(r"(\*\*.*?\*\*)", para_text)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            run = p.add_run(part[2:-2])
                            run.bold = True
                        else:
                            p.add_run(part)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.docx"'},
        )

    # ── CSV ────────────────────────────────────────────────
    if format == "csv":
        out = io.StringIO()
        writer = csv_mod.writer(out)

        # Try to find tables in results; fallback to Node/Content columns
        has_tables = False
        for nid, content in results.items():
            tables = _parse_markdown_tables(str(content))
            if tables:
                has_tables = True
                label = node_labels.get(nid, nid)
                writer.writerow([f"--- {label} ---"])
                for row in tables:
                    writer.writerow(row)
                writer.writerow([])

        if not has_tables:
            writer.writerow(["Node", "Content"])
            for nid, content in results.items():
                label = node_labels.get(nid, nid)
                writer.writerow([label, _md_to_plain(str(content))])

        buf = io.BytesIO(out.getvalue().encode("utf-8-sig"))  # BOM for Excel compat
        return StreamingResponse(
            buf,
            media_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.csv"'},
        )

    # ── XLSX ───────────────────────────────────────────────
    if format == "xlsx":
        from openpyxl import Workbook
        from openpyxl.styles import Font

        wb = Workbook()
        wb.remove(wb.active)  # Remove default sheet

        for nid, content in results.items():
            label = node_labels.get(nid, nid)
            # Sheet name max 31 chars
            sheet_name = label[:31] if label else nid[:31]
            ws = wb.create_sheet(title=sheet_name)

            tables = _parse_markdown_tables(str(content))
            if tables:
                for i, row in enumerate(tables):
                    for j, cell in enumerate(row):
                        c = ws.cell(row=i + 1, column=j + 1, value=cell)
                        if i == 0:
                            c.font = Font(bold=True)
            else:
                # Write text content line-by-line
                ws.cell(row=1, column=1, value=label).font = Font(bold=True)
                plain = _md_to_plain(str(content))
                for i, line in enumerate(plain.splitlines(), 2):
                    ws.cell(row=i, column=1, value=line)
                # Auto-width for column A
                ws.column_dimensions["A"].width = 80

        if not wb.sheetnames:
            wb.create_sheet(title="Results")

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.xlsx"'},
        )

    # ── PNG ────────────────────────────────────────────────
    if format == "png":
        from fpdf import FPDF

        # Generate PDF first (reuse PDF logic)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        usable_w = pdf.w - pdf.l_margin - pdf.r_margin

        def _safe_png(text: str) -> str:
            return text.encode("latin-1", "replace").decode("latin-1")

        pdf.set_font("Helvetica", "B", 16)
        pdf.cell(usable_w, 10, _safe_png(wf.name), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        for nid, content in results.items():
            label = node_labels.get(nid, nid)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_x(pdf.l_margin)
            pdf.cell(usable_w, 8, _safe_png(label), new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 10)
            plain = _md_to_plain(str(content))
            for line in plain.splitlines():
                if not line.strip():
                    pdf.ln(3)
                else:
                    pdf.set_x(pdf.l_margin)
                    pdf.multi_cell(usable_w, 5, _safe_png(line))
            pdf.ln(6)

        pdf_bytes = pdf.output()

        # Convert PDF pages to PNG using PyMuPDF
        try:
            import fitz  # PyMuPDF

            doc_pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
            images = []
            dpi = 150
            matrix = fitz.Matrix(dpi / 72, dpi / 72)

            for page in doc_pdf:
                pix = page.get_pixmap(matrix=matrix)
                img_bytes = pix.tobytes("png")
                images.append(img_bytes)
            doc_pdf.close()

            if len(images) == 1:
                buf = io.BytesIO(images[0])
            else:
                # Stack pages vertically using Pillow
                from PIL import Image as PILImage
                pil_images = [PILImage.open(io.BytesIO(b)) for b in images]
                total_height = sum(im.height for im in pil_images)
                max_width = max(im.width for im in pil_images)
                combined = PILImage.new("RGB", (max_width, total_height), (255, 255, 255))
                y_offset = 0
                for im in pil_images:
                    combined.paste(im, (0, y_offset))
                    y_offset += im.height
                buf = io.BytesIO()
                combined.save(buf, format="PNG")
                buf.seek(0)

            return StreamingResponse(
                buf,
                media_type="image/png",
                headers={"Content-Disposition": f'attachment; filename="{wf.name}.png"'},
            )
        except ImportError:
            raise HTTPException(500, "PyMuPDF (fitz) not installed. Run: pip install PyMuPDF")

    # ── GeoJSON ─────────────────────────────────────────────
    if format == "geojson":
        import re as _re
        geojson_pattern = _re.compile(r'\{"type"\s*:\s*"FeatureCollection"[\s\S]*?\}(?:\s*\]?\s*\})')
        features = []
        for content in results.values():
            # Try to find GeoJSON in artifact blocks
            artifact_re = _re.compile(r'```artifact\s*\n([\s\S]*?)```')
            for match in artifact_re.finditer(str(content)):
                try:
                    obj = json.loads(match.group(1).strip())
                    if obj.get("type") == "geojson" and obj.get("data"):
                        data = obj["data"] if isinstance(obj["data"], dict) else json.loads(obj["data"])
                        if data.get("type") == "FeatureCollection":
                            features.extend(data.get("features", []))
                except (json.JSONDecodeError, TypeError):
                    pass
            # Also try raw GeoJSON in content
            for match in geojson_pattern.finditer(str(content)):
                try:
                    data = json.loads(match.group())
                    if data.get("type") == "FeatureCollection":
                        features.extend(data.get("features", []))
                except (json.JSONDecodeError, TypeError):
                    pass

        if not features:
            # Fallback: wrap text results as GeoJSON with no geometry
            raise HTTPException(404, "No GeoJSON/spatial data found in workflow results")

        fc = {"type": "FeatureCollection", "features": features}
        geojson_bytes = json.dumps(fc, ensure_ascii=False, indent=2).encode("utf-8")
        return StreamingResponse(
            io.BytesIO(geojson_bytes),
            media_type="application/geo+json",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}.geojson"'},
        )

    # ── Shapefile (as ZIP) ─────────────────────────────────
    if format == "shapefile":
        try:
            import geopandas as gpd
            from shapely.geometry import shape
        except ImportError:
            raise HTTPException(500, "geopandas/shapely not installed")

        import re as _re
        features = []
        artifact_re = _re.compile(r'```artifact\s*\n([\s\S]*?)```')
        for content in results.values():
            for match in artifact_re.finditer(str(content)):
                try:
                    obj = json.loads(match.group(1).strip())
                    if obj.get("type") == "geojson" and obj.get("data"):
                        data = obj["data"] if isinstance(obj["data"], dict) else json.loads(obj["data"])
                        if data.get("type") == "FeatureCollection":
                            features.extend(data.get("features", []))
                except (json.JSONDecodeError, TypeError):
                    pass

        if not features:
            raise HTTPException(404, "No spatial data found in workflow results for shapefile export")

        gdf = gpd.GeoDataFrame.from_features(features, crs="EPSG:4326")
        import tempfile
        with tempfile.TemporaryDirectory() as tmpdir:
            shp_path = os.path.join(tmpdir, f"{wf.name}.shp")
            gdf.to_file(shp_path)
            buf = io.BytesIO()
            with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
                for f in os.listdir(tmpdir):
                    zf.write(os.path.join(tmpdir, f), f)
            buf.seek(0)

        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}_shp.zip"'},
        )

    # ── ZIP All (every format bundled) ─────────────────────
    if format == "zip_all":
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            # Markdown files per node
            for i, (nid, content) in enumerate(results.items(), 1):
                label = node_labels.get(nid, nid).replace(" ", "_")
                zf.writestr(f"markdown/{i:02d}_{label}.md", str(content))

            # Combined markdown
            md_lines = [f"# {wf.name} — Risultati\n"]
            for nid, content in results.items():
                label = node_labels.get(nid, nid)
                md_lines.append(f"## {label}\n\n{content}\n")
            zf.writestr(f"{wf.name}.md", "\n".join(md_lines))

            # CSV
            csv_buf = io.StringIO()
            csv_buf.write("\ufeff")  # BOM
            import csv as csv_mod
            writer = csv_mod.writer(csv_buf)
            writer.writerow(["Node", "Content"])
            for nid, content in results.items():
                label = node_labels.get(nid, nid)
                writer.writerow([label, str(content)[:32000]])
            zf.writestr(f"{wf.name}.csv", csv_buf.getvalue())

            # PDF
            try:
                from fpdf import FPDF
                pdf = FPDF()
                pdf.set_auto_page_break(auto=True, margin=15)
                for nid, content in results.items():
                    label = node_labels.get(nid, nid)
                    pdf.add_page()
                    pdf.set_font("Helvetica", "B", 14)
                    pdf.cell(0, 10, label, new_x="LMARGIN", new_y="NEXT")
                    pdf.set_font("Helvetica", size=9)
                    text = _md_to_plain(str(content))
                    pdf.multi_cell(0, 5, text)
                zf.writestr(f"{wf.name}.pdf", pdf.output())
            except ImportError:
                pass

            # GeoJSON (if spatial data present)
            import re as _re
            features = []
            artifact_re = _re.compile(r'```artifact\s*\n([\s\S]*?)```')
            for content in results.values():
                for match in artifact_re.finditer(str(content)):
                    try:
                        obj = json.loads(match.group(1).strip())
                        if obj.get("type") == "geojson" and obj.get("data"):
                            data = obj["data"] if isinstance(obj["data"], dict) else json.loads(obj["data"])
                            if data.get("type") == "FeatureCollection":
                                features.extend(data.get("features", []))
                    except (json.JSONDecodeError, TypeError):
                        pass
            if features:
                fc = {"type": "FeatureCollection", "features": features}
                zf.writestr(f"{wf.name}.geojson", json.dumps(fc, ensure_ascii=False, indent=2))

        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/zip",
            headers={"Content-Disposition": f'attachment; filename="{wf.name}_all.zip"'},
        )

    # ── ZIP (default) ─────────────────────────────────────
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for i, (nid, content) in enumerate(results.items(), 1):
            label = node_labels.get(nid, nid).replace(" ", "_")
            zf.writestr(f"{i:02d}_{label}.md", str(content))
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{wf.name}.zip"'},
    )


# ── Templates & Node Registry ──────────────────────────────


@router.get("/templates")
async def list_templates():
    """Return all predefined workflow templates."""
    from builder.templates import get_all_templates
    return get_all_templates()


@router.get("/templates/{template_id}")
async def get_template(template_id: str):
    """Return a single template by ID."""
    from builder.templates import get_template_by_id
    tpl = get_template_by_id(template_id)
    if not tpl:
        raise HTTPException(404, "Template not found")
    return tpl


@router.post("/templates/{template_id}/instantiate", status_code=201)
async def instantiate_template(template_id: str, db: AsyncSession = Depends(get_db)):
    """Create a new workflow from a template."""
    from builder.templates import get_template_by_id
    tpl = get_template_by_id(template_id)
    if not tpl:
        raise HTTPException(404, "Template not found")
    row = WorkflowRow(
        id=str(uuid.uuid4()),
        name=tpl["name"],
        description=tpl["description"],
        definition=tpl["definition"],
    )
    db.add(row)
    await db.commit()
    return WorkflowOut(
        id=row.id, name=row.name, description=row.description,
        definition=tpl["definition"],
        created_at=row.created_at.isoformat(),
        updated_at=row.created_at.isoformat(),
    )


@router.get("/node-registry")
async def node_registry():
    """Return the available node types for the workflow builder."""
    from builder.node_registry import get_registry
    return get_registry()
